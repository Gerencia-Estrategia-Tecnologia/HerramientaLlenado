import pandas as pd
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Definición de las dimensiones y categorías
dimensiones = ["Mercado", "Servicio", "Producto", "Datos", "Tecnología", "Aliados", "Capacidades"]
categorias = ["Núcleo integrado", "Servicios digitales", "Marketplace digital", "Ecosistema digital"]

# Inicialización de un diccionario para almacenar las respuestas
matriz = {dim: {cat: "" for cat in categorias} for dim in dimensiones}

# Iteración para pedir al usuario que complete la matriz
for dim in dimensiones:
    for cat in categorias:
        respuesta = input(f"Ingrese el contenido para '{dim}' en la categoría '{cat}': ")
        matriz[dim][cat] = respuesta

# Creación de un DataFrame de pandas
df = pd.DataFrame(matriz).T

# Crear un nuevo archivo de Excel con openpyxl
wb = Workbook()
ws = wb.active
ws.title = "Matriz Completada"

# Añadir las filas del DataFrame al worksheet
for r_idx, row in enumerate(dataframe_to_rows(df, index=True, header=True), 1):
    for c_idx, value in enumerate(row, 1):
        ws.cell(row=r_idx, column=c_idx, value=value)

# Aplicar estilos
header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
header_font = Font(color="FFFFFF", bold=True)
alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

# Aplicar estilo a los encabezados de las categorías (primera fila)
for cell in ws[1]:
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = alignment

# Aplicar estilo a los encabezados de las dimensiones (primera columna)
for cell in ws['A']:
    if cell.row != 1:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = alignment

# Ajuste de ancho de columnas y alto de filas
for col in ws.columns:
    max_length = 0
    column = col[0].column_letter
    for cell in col:
        try:
            if len(str(cell.value)) > max_length:
                max_length = len(cell.value)
        except:
            pass
    adjusted_width = (max_length + 2)
    ws.column_dimensions[column].width = adjusted_width

# Aplicar bordes a todas las celdas
thin_border = Border(left=Side(style='thin'),
                     right=Side(style='thin'),
                     top=Side(style='thin'),
                     bottom=Side(style='thin'))

for row in ws.iter_rows():
    for cell in row:
        cell.border = thin_border

# Guardar el archivo de Excel
output_file = "matriz_completada_formateada.xlsx"
wb.save(output_file)

# Generar el informe en Word
doc = Document()
doc.add_heading('Encuentro binacional – Informe Final', 0)

# Agregar el espacio para comentarios generales
doc.add_paragraph("Aquí se deja este espacio para los comentarios…", style='Normal')

# Insertar título "Formato Final" con tamaño de fuente 16
format_final_heading = doc.add_heading('Formato Final', level=1)
format_final_heading.runs[0].font.size = Pt(16)
format_final_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Cambiar color a negro

# Insertar la tabla con los datos de la matriz en el documento debajo de "Formato Final" y antes de los comentarios
table = doc.add_table(rows=len(df)+1, cols=len(df.columns)+1)

# Llenar los encabezados de la tabla
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Dimensiones'
for i, cat in enumerate(categorias):
    hdr_cells[i+1].text = cat

# Aplicar formato a los encabezados (color de fondo y texto)
for cell in table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Texto blanco
    cell.paragraphs[0].runs[0].font.bold = True
    tc_pr = cell._element.get_or_add_tcPr()
    cell_fill = OxmlElement('w:shd')
    cell_fill.set(qn('w:fill'), '1F4E78')  # Fondo oscuro
    tc_pr.append(cell_fill)

# Llenar el contenido de la tabla con las dimensiones y categorías
for i, dim in enumerate(dimensiones):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = dim
    for j, cat in enumerate(categorias):
        row_cells[j+1].text = matriz[dim][cat] if matriz[dim][cat] else ""

# Aplicar bordes y centrar el texto
for row in table.rows:
    for cell in row.cells:
        cell.paragraphs[0].alignment = 1  # Centrar el texto
        tc_pr = cell._element.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Tamaño del borde
            tc_borders.append(border)
        tc_pr.append(tc_borders)

# Aplicar tamaño de fuente a las celdas
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)  # Tamaño de letra 12

# Agregar el espacio para comentarios adicionales después de la tabla
doc.add_paragraph("Aquí se deja este espacio para los comentarios…", style='Normal')

# Agregar el contenido general con tamaño de fuente 16
doc.add_heading('Contenido General', level=1)
content_general_heading = doc.paragraphs[-1]
content_general_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Cambiar color a negro
content_general_heading.runs[0].font.size = Pt(16)
doc.add_paragraph("Aquí se deja este espacio para los comentarios…", style='Normal')

# Iterar sobre las dimensiones y categorías
for dim in dimensiones:
    heading_dim = doc.add_heading(dim, level=2)
    heading_dim.runs[0].font.size = Pt(13)
    heading_dim.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Cambiar color a negro
    heading_dim.paragraph_format.left_indent = Pt(10)  # Sangría izquierda

    for cat in categorias:
        contenido = matriz[dim][cat]
        heading = doc.add_heading(f'{cat}:', level=3)
        heading.runs[0].font.size = Pt(13)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Cambiar color a negro
        heading.paragraph_format.left_indent = Pt(10)

        p = doc.add_paragraph("Aquí se deja este espacio para los comentarios…", style='Normal')
        p.paragraph_format.left_indent = Pt(10)

        if contenido:
            elementos = [elem.strip() for elem in contenido.split(',') if elem.strip()]
            for elem in elementos:
                p = doc.add_paragraph(elem, style='ListBullet')
                p.paragraph_format.left_indent = Pt(50)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)

# Aplicar un mayor interlineado a todo el documento
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.space_before = Pt(12)

# Guardar el documento Word
output_word = "informe_matriz.docx"
doc.save(output_word)

print(f"¡Matriz completada y formateada! El archivo Excel '{output_file}' y el informe Word '{output_word}' se han guardado exitosamente.")
